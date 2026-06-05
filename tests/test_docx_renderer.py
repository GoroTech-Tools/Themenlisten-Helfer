import datetime
import os
import sys
import tempfile
import unittest

from docx import Document
import pandas as pd


SRC_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src'))
if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

from docx_renderer import aktueller_stand_text, erstelle_dateiname, baue_ersetzungen, render_vorlage


class TestDocxRenderer(unittest.TestCase):
    def test_aktueller_stand_text(self):
        self.assertEqual(aktueller_stand_text(datetime.date(2026, 6, 5)), 'Juni 2026')

    def test_erstelle_dateiname(self):
        name = erstelle_dateiname('LB 1', 'Müller', 'Juni 2026')
        self.assertEqual(name, 'Themenliste_LB_1_Müller_Juni_2026.docx')

    def test_render_vorlage_und_ersetzungen(self):
        row = pd.Series(
            {
                'Vorname': 'Anna',
                'Nachname': 'Beispiel',
                'Anrede': 'Frau',
                'Zeitraum': 'Juni 2026',
                'Kurs': 'AP1',
            }
        )

        with tempfile.TemporaryDirectory() as tmp:
            vorlage = os.path.join(tmp, 'template.docx')
            ziel = os.path.join(tmp, 'output.docx')

            doc = Document()
            doc.add_paragraph('Hallo <<vorname>> <<nachname>>')
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = 'Bereich: <<lernbereich>>'
            doc.sections[0].header.paragraphs[0].text = 'Stand: <<stand>>'
            doc.sections[0].footer.paragraphs[0].text = 'Datei: <<dateiname>>'
            doc.save(vorlage)

            dateiname = erstelle_dateiname('LB1', 'Beispiel', 'Juni 2026')
            ersetzungen = baue_ersetzungen(row, 'LB1', dateiname, stand='Juni 2026')
            render_vorlage(vorlage, ziel, ersetzungen)

            rendered = Document(ziel)
            self.assertIn('Hallo Anna Beispiel', rendered.paragraphs[0].text)
            self.assertIn('Bereich: LB1', rendered.tables[0].cell(0, 0).text)
            self.assertIn('Stand: Juni 2026', rendered.sections[0].header.paragraphs[0].text)
            self.assertIn(f'Datei: {dateiname}', rendered.sections[0].footer.paragraphs[0].text)


if __name__ == '__main__':
    unittest.main()
