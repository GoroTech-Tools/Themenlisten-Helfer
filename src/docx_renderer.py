import datetime
from typing import Any, Mapping

from docx import Document


def aktueller_stand_text(heute: datetime.date | None = None) -> str:
    monate = [
        'Januar', 'Februar', 'März', 'April', 'Mai', 'Juni',
        'Juli', 'August', 'September', 'Oktober', 'November', 'Dezember'
    ]
    heute = heute or datetime.date.today()
    return f"{monate[heute.month - 1]} {heute.year}"


def erstelle_dateiname(lernbereich: str, nachname: str, zeitraum: str) -> str:
    return f"Themenliste_{lernbereich}_{nachname}_{zeitraum}.docx".replace(' ', '_')


def baue_ersetzungen(row: Any, lernbereich: str, dateiname: str, stand: str | None = None) -> dict[str, str]:
    return {
        '<<vorname>>': str(row['Vorname']),
        '<<nachname>>': str(row['Nachname']),
        '<<anrede>>': str(row['Anrede']),
        '<<zeitraum>>': str(row['Zeitraum']),
        '<<kurs>>': str(row['Kurs']),
        '<<lernbereich>>': str(lernbereich),
        '<<dateiname>>': str(dateiname),
        '<<stand>>': str(stand or aktueller_stand_text()),
    }


def _ersetze_platzhalter_in_paragraphs(paragraphs: Any, ersetzungen: Mapping[str, str]) -> None:
    for p in paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        replaced = full_text
        for key, value in ersetzungen.items():
            replaced = replaced.replace(key, value)
        if replaced != full_text:
            for i in range(len(p.runs) - 1, -1, -1):
                p.runs[i].clear()
                p.runs[i].text = ''
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)


def _ersetze_platzhalter_in_table(table: Any, ersetzungen: Mapping[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            _ersetze_platzhalter_in_paragraphs(cell.paragraphs, ersetzungen)
            for nested in cell.tables:
                _ersetze_platzhalter_in_table(nested, ersetzungen)


def render_vorlage(vorlage: str, ziel: str, ersetzungen: Mapping[str, str]) -> None:
    doc = Document(vorlage)

    for section in doc.sections:
        _ersetze_platzhalter_in_paragraphs(section.header.paragraphs, ersetzungen)
        for table in section.header.tables:
            _ersetze_platzhalter_in_table(table, ersetzungen)
        _ersetze_platzhalter_in_paragraphs(section.footer.paragraphs, ersetzungen)
        for table in section.footer.tables:
            _ersetze_platzhalter_in_table(table, ersetzungen)

    _ersetze_platzhalter_in_paragraphs(doc.paragraphs, ersetzungen)
    for table in doc.tables:
        _ersetze_platzhalter_in_table(table, ersetzungen)

    doc.save(ziel)
