from PIL import Image, ImageTk  # Pillow muss installiert sein: pip install pillow
import sys
import os
from typing import Optional

# === Basisverzeichnis bestimmen (EXE-kompatibel) ===
def get_app_dir() -> str:
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def get_resource_dirs() -> list[str]:
    dirs: list[str] = []
    meipass = getattr(sys, '_MEIPASS', None)
    if getattr(sys, 'frozen', False) and isinstance(meipass, str) and meipass:
        dirs.append(meipass)
    dirs.append(get_app_dir())
    return dirs

def resolve_path(*candidates: str, must_exist: bool = False) -> Optional[str]:
    for path in candidates:
        if must_exist:
            if path and os.path.exists(path):
                return path
        else:
            if path:
                return path
    return None
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
import pandas as pd
import glob
from docx import Document
import win32com.client as win32
import re
import time
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import threading
import subprocess
import pythoncom

# === Hilfsfunktionen ===
def ersetze_umlaute(s):
    s = re.sub(r'[äÄ]', 'ae', s)
    s = re.sub(r'[öÖ]', 'oe', s)
    s = re.sub(r'[üÜ]', 'ue', s)
    s = re.sub(r'ß', 'ss', s)
    return s

def finde_vorlagen(vorlagen_ordner, lernbereich=None):
    endungen = ('.docx', '.dotm')
    treffer = []
    for endung in endungen:
        if lernbereich is None:
            muster = os.path.join(vorlagen_ordner, f'*{endung}')
        else:
            muster = os.path.join(vorlagen_ordner, f'Themenliste_{lernbereich}*{endung}')
        treffer.extend(glob.glob(muster))
    return sorted(set(treffer))

def erstelle_themenlisten(excel_path, status_callback):
    result = {
        'selected': 0,
        'created': 0,
        'missing_templates': 0,
        'aborted': False
    }
    df = pd.read_excel(excel_path, sheet_name='Teilnehmende')
    df = df[df['Verarbeiten'].str.lower() == 'ja']
    result['selected'] = len(df)
    app_dir = get_app_dir()
    resource_dirs = get_resource_dirs()
    vorlagen_ordner = resolve_path(
        *[os.path.join(d, 'data', 'Themenlisten-Vorlagen') for d in resource_dirs],
        *[os.path.join(d, 'templates', 'Themenlisten-Vorlagen') for d in resource_dirs],
        *[os.path.join(d, 'Themenlisten-Vorlagen') for d in resource_dirs],
        must_exist=True
    )
    if not vorlagen_ordner:
        status_callback("[FEHLER] Kein Vorlagenordner gefunden (erwartet: data/Themenlisten-Vorlagen, templates/Themenlisten-Vorlagen oder Themenlisten-Vorlagen).")
        result['aborted'] = True
        return result
    vorhandene_vorlagen = finde_vorlagen(vorlagen_ordner)
    if not vorhandene_vorlagen:
        status_callback(f"[FEHLER] Im Vorlagenordner wurden keine .docx- oder .dotm-Dateien gefunden: {vorlagen_ordner}")
        result['aborted'] = True
        return result
    ausgabe_ordner = os.path.join(app_dir, 'Themenlisten')
    os.makedirs(ausgabe_ordner, exist_ok=True)
    for idx, row in df.iterrows():
        lernbereich = str(row['Lernbereich'])
        vorlagen = finde_vorlagen(vorlagen_ordner, lernbereich)
        if not vorlagen:
            status_callback(f"[WARNUNG] Keine Vorlage für Lernbereich '{lernbereich}' gefunden!")
            result['missing_templates'] += 1
            continue
        vorlage = vorlagen[0]
        # Platzhalter <<stand>>: aktueller Monat ausgeschrieben und Jahr
        import datetime
        monate = [
            'Januar', 'Februar', 'März', 'April', 'Mai', 'Juni',
            'Juli', 'August', 'September', 'Oktober', 'November', 'Dezember'
        ]
        heute = datetime.date.today()
        stand = f"{monate[heute.month-1]} {heute.year}"
        dateiname = f"Themenliste_{lernbereich}_{row['Nachname']}_{row['Zeitraum']}.docx".replace(' ', '_')
        ersetzungen = {
            '<<vorname>>': str(row['Vorname']),
            '<<nachname>>': str(row['Nachname']),
            '<<anrede>>': str(row['Anrede']),
            '<<zeitraum>>': str(row['Zeitraum']),
            '<<kurs>>': str(row['Kurs']),
            '<<lernbereich>>': lernbereich,
            '<<dateiname>>': dateiname,
            '<<stand>>': stand
        }
        doc = Document(vorlage)
        def ersetze_platzhalter_in_paragraphs(paragraphs):
            for p in paragraphs:
                full_text = ''.join(run.text for run in p.runs)
                replaced = full_text
                for key, value in ersetzungen.items():
                    replaced = replaced.replace(key, value)
                if replaced != full_text:
                    for i in range(len(p.runs)-1, -1, -1):
                        p.runs[i].clear()
                        p.runs[i].text = ''
                    if p.runs:
                        p.runs[0].text = replaced
                    else:
                        p.add_run(replaced)

        def ersetze_platzhalter_in_table(table):
            for row in table.rows:
                for cell in row.cells:
                    ersetze_platzhalter_in_paragraphs(cell.paragraphs)
                    for t in cell.tables:
                        ersetze_platzhalter_in_table(t)

        for section in doc.sections:
            ersetze_platzhalter_in_paragraphs(section.header.paragraphs)
            for table in section.header.tables:
                ersetze_platzhalter_in_table(table)
            ersetze_platzhalter_in_paragraphs(section.footer.paragraphs)
            for table in section.footer.tables:
                ersetze_platzhalter_in_table(table)
        ersetze_platzhalter_in_paragraphs(doc.paragraphs)
        for table in doc.tables:
            ersetze_platzhalter_in_table(table)
        dateiname = f"Themenliste_{lernbereich}_{row['Nachname']}_{row['Zeitraum']}.docx".replace(' ', '_')
        ziel = os.path.join(ausgabe_ordner, dateiname)
        doc.save(ziel)
        result['created'] += 1
        status_callback(f"Erstellt: {ziel}")
    status_callback(f"FERTIG: Themenlisten erstellt ({result['created']} von {result['selected']}).")
    return result

def erstelle_emails(excel_path, status_callback):
    result = {
        'selected': 0,
        'drafts': 0,
        'missing_templates': 0,
        'missing_attachments': 0,
        'missing_mail_config': 0,
        'aborted': False
    }
    personen = pd.read_excel(excel_path, sheet_name='Teilnehmende')
    personen = personen[personen['Verarbeiten'].str.lower() == 'ja']
    result['selected'] = len(personen)
    mailkonf = pd.read_excel(excel_path, sheet_name='E-Mail-Konfiguration')
    app_dir = get_app_dir()
    resource_dirs = get_resource_dirs()
    vorlagen_ordner = resolve_path(
        *[os.path.join(d, 'data', 'Themenlisten-Vorlagen') for d in resource_dirs],
        *[os.path.join(d, 'templates', 'Themenlisten-Vorlagen') for d in resource_dirs],
        *[os.path.join(d, 'Themenlisten-Vorlagen') for d in resource_dirs],
        must_exist=True
    )
    if not vorlagen_ordner:
        status_callback("[FEHLER] Kein Vorlagenordner gefunden (erwartet: data/Themenlisten-Vorlagen, templates/Themenlisten-Vorlagen oder Themenlisten-Vorlagen).")
        result['aborted'] = True
        return result
    vorhandene_vorlagen = finde_vorlagen(vorlagen_ordner)
    if not vorhandene_vorlagen:
        status_callback(f"[FEHLER] Im Vorlagenordner wurden keine .docx- oder .dotm-Dateien gefunden: {vorlagen_ordner}")
        result['aborted'] = True
        return result
    ausgabe_ordner = os.path.join(app_dir, 'Themenlisten')
    os.makedirs(ausgabe_ordner, exist_ok=True)
    outlook = win32.Dispatch('Outlook.Application')
    for idx, row in personen.iterrows():
        lernbereich = str(row['Lernbereich'])
        vorlagen = finde_vorlagen(vorlagen_ordner, lernbereich)
        if not vorlagen:
            status_callback(f"[WARNUNG] Keine Vorlage für Lernbereich '{lernbereich}' gefunden!")
            result['missing_templates'] += 1
            continue
        vorlage = vorlagen[0]
        dateiname = f"Themenliste_{lernbereich}_{row['Nachname']}_{row['Zeitraum']}.docx".replace(' ', '_')
        anhang = os.path.join(ausgabe_ordner, dateiname)
        if not os.path.exists(anhang):
            status_callback(f"[WARNUNG] Anhang nicht gefunden: {anhang}")
            result['missing_attachments'] += 1
            continue
        # E-Mail-Regeln: mehrere Vornamen zusammenziehen, ggf. nur erster Buchstabe
        vornamen_roh = str(row['Vorname']).lower()
        nachname = ersetze_umlaute(str(row['Nachname']).lower())
        vornamen_ohne_leer = ersetze_umlaute(vornamen_roh.replace(' ', ''))
        empfaenger_vorname = vornamen_ohne_leer
        if len(empfaenger_vorname + '.' + nachname) > 20:
            empfaenger_vorname = ersetze_umlaute(vornamen_roh[0])
        empfaenger = f"{empfaenger_vorname}.{nachname}@ab.bfw.local"
        template_bez = os.path.splitext(os.path.basename(vorlage))[0]
        mailrow = mailkonf[mailkonf['TemplateBezeichnung'] == template_bez]
        if mailrow.empty:
            status_callback(f"[WARNUNG] Keine E-Mail-Konfiguration für Vorlage {template_bez}")
            result['missing_mail_config'] += 1
            continue
        subject = str(mailrow.iloc[0]['Subject'])
        body = str(mailrow.iloc[0]['BodyText']).replace('\\n', '\n')
        for key, value in {
            '<<vorname>>': row['Vorname'],
            '<<nachname>>': row['Nachname'],
            '<<anrede>>': row['Anrede'],
            '<<zeitraum>>': row['Zeitraum'],
            '<<kurs>>': row['Kurs'],
            '<<lernbereich>>': row['Lernbereich']
        }.items():
            subject = subject.replace(key, str(value))
            body = body.replace(key, str(value))
        mail = outlook.CreateItem(0)
        mail.To = empfaenger
        mail.Subject = subject
        mail.Body = body
        mail.Attachments.Add(anhang)
        mail.Display()
        result['drafts'] += 1
        status_callback(f"E-Mail-Entwurf für {empfaenger} angezeigt.")
    status_callback(f"FERTIG: E-Mail-Entwürfe erstellt ({result['drafts']} von {result['selected']}).")
    return result

def loesche_alte_themenlisten(excel_path, status_callback):
    app_dir = get_app_dir()
    ausgabe_ordner = os.path.join(app_dir, 'Themenlisten')
    os.makedirs(ausgabe_ordner, exist_ok=True)
    jetzt = time.time()
    ein_tag = 24 * 60 * 60
    geloescht = 0
    for fname in os.listdir(ausgabe_ordner):
        pfad = os.path.join(ausgabe_ordner, fname)
        if os.path.isfile(pfad):
            # TXT-Dateien niemals löschen
            if fname.lower().endswith('.txt'):
                continue
            mtime = os.path.getmtime(pfad)
            if jetzt - mtime > ein_tag:
                try:
                    os.remove(pfad)
                    status_callback(f"Gelöscht: {pfad}")
                    geloescht += 1
                except Exception as e:
                    status_callback(f"[FEHLER] Konnte {pfad} nicht löschen: {e}")
    status_callback(f"Alte Themenlisten gelöscht: {geloescht}")
    return geloescht

# === GUI ===
class ThemenlistenApp:
    def __init__(self, root):
        self.root = root
        # Status-Text für Fehlermeldungen und Hinweise
        self.status_text = tk.StringVar()
        # Versionsnummer aus version.txt lesen
        try:
            resource_dirs = get_resource_dirs()
            version_path = resolve_path(
                *[os.path.join(d, "src", "version.txt") for d in resource_dirs],
                *[os.path.join(d, "config", "version.txt") for d in resource_dirs],
                *[os.path.join(d, "version.txt") for d in resource_dirs],
                must_exist=True
            )
            if version_path:
                with open(version_path, "r", encoding="utf-8") as vf:
                    version = vf.read().strip()
            else:
                version = ""
        except Exception:
            version = ""
        self.root.title(f"Themenlistenhelfer {version}")
        self.root.geometry("800x400")  # Noch breiter und höher für Bild
        self.excel_path = tk.StringVar()
        # Fortschrittsbalken statt Status-Text
        self.progress = tk.DoubleVar()
        # Standarddatei (immer im EXE-Verzeichnis suchen)
        app_dir = get_app_dir()
        resource_dirs = get_resource_dirs()
        default_excel = resolve_path(
            os.path.join(app_dir, 'data', 'Auswahl Teilnehmende zu Lernbereichen.xlsx'),
            os.path.join(app_dir, 'Auswahl Teilnehmende zu Lernbereichen.xlsx'),
            *[os.path.join(d, 'data', 'Auswahl Teilnehmende zu Lernbereichen.xlsx') for d in resource_dirs],
            *[os.path.join(d, 'Auswahl Teilnehmende zu Lernbereichen.xlsx') for d in resource_dirs],
            must_exist=True
        )
        if not default_excel:
            default_excel = os.path.join(app_dir, 'data', 'Auswahl Teilnehmende zu Lernbereichen.xlsx')
        self.excel_path.set(default_excel)
        frame_outer = tk.Frame(root, padx=20, pady=10)
        frame_outer.pack(fill=tk.BOTH, expand=True)
        # Status-Label ganz oben
        status_label = tk.Label(frame_outer, textvariable=self.status_text, fg="red", anchor="w", font=("Calibri", 10))
        status_label.pack(fill=tk.X, pady=(0, 5))

        # Haupt-Layout: 2 Spalten (links: Controls, rechts: Anleitung)
        frame_main = tk.Frame(frame_outer)
        frame_main.pack(fill=tk.BOTH, expand=True)

        # Linke Spalte: Label und Buttons
        frame_left = tk.Frame(frame_main)
        frame_left.grid(row=0, column=0, sticky="nw")
        label_excel = tk.Label(frame_left, text="Excel-Datei: Auswahl Teilnehmende zu Lernbereichen.xlsx", anchor="w", justify="left")
        try:
            label_excel.configure(font=("Calibri", 11))
        except Exception:
            pass
        label_excel.pack(anchor="w", pady=(0,5))
        btn_excel = tk.Button(frame_left, text="Auswahl Teilnehmende zu Lernbereichen", command=self.excel_bearbeiten, bg="#FFA500", fg="black", activebackground="#FFB733", activeforeground="white")
        btn_start = tk.Button(frame_left, text="Themenlistenhelfer starten", command=self.batch_starten, bg="#4CAF50", fg="white", activebackground="#6FD36F", activeforeground="white")
        try:
            btn_excel.configure(font=("Calibri", 11))
            btn_start.configure(font=("Calibri", 11))
        except Exception:
            pass
        btn_excel.pack(anchor="w", pady=2, fill=tk.X)
        btn_start.pack(anchor="w", pady=2, fill=tk.X)

        # Rechte Spalte: Anleitung als Textbox (read-only)
        frame_right = tk.Frame(frame_main)
        frame_right.grid(row=0, column=1, sticky="ne", padx=(30,0))
        anleitung_fett = "Anleitung:"
        anleitung_rest = (
            "\n\n1. Excel-Datei anpassen (z. B. Teilnehmende, Lernbereiche, E-Mail-Konfiguration).\n"
            "2. 'Themenlistenhelfer starten' ausführen.\n"
            "\nDie generierten Themenlisten und E-Mail-Entwürfe werden automatisch erstellt."
        )
        textbox = tk.Text(frame_right, width=48, height=10, wrap="word", bg=frame_outer.cget("bg"), relief="flat", borderwidth=0)
        try:
            textbox.configure(font=("Calibri", 11))
        except Exception:
            pass  # Falls Calibri nicht verfügbar ist, Standard verwenden
        textbox.tag_configure("fett", font=("Calibri", 11, "bold"))
        textbox.insert("1.0", anleitung_fett, "fett")
        textbox.insert("1.end", anleitung_rest)
        textbox.config(state="disabled")
        textbox.pack(anchor="n", pady=5)

        # Fortschrittsbalken unten (grau)
        style = ttk.Style()
        style.theme_use('default')
        style.configure("grey.Horizontal.TProgressbar", troughcolor="#e0e0e0", background="#888888", bordercolor="#e0e0e0", lightcolor="#bbbbbb", darkcolor="#888888")
        self.progressbar = ttk.Progressbar(frame_outer, variable=self.progress, maximum=100, style="grey.Horizontal.TProgressbar")
        self.progressbar.pack(pady=10, fill=tk.X)
        # Label für Fertig-Anzeige
        self.progress_label = tk.Label(frame_outer, text="", font=("Calibri", 11), fg="#4CAF50")
        self.progress_label.pack()

        # Bild unten rechts anzeigen, falls vorhanden, mit Fehlerausgabe
        try:
            resource_dirs = get_resource_dirs()
            img_path = resolve_path(
                *[os.path.join(d, "assets", "images", "Themenlistenhelfer.png") for d in resource_dirs],
                *[os.path.join(d, "Themenlistenhelfer.png") for d in resource_dirs],
                must_exist=True
            )
            if img_path and os.path.exists(img_path):
                img = Image.open(img_path)
                img = img.resize((64, 64), Image.Resampling.LANCZOS)
                self.tk_img = ImageTk.PhotoImage(img)
                img_label = tk.Label(frame_outer, image=self.tk_img)
                img_label.pack(side=tk.RIGHT, anchor="se", padx=10, pady=10)
            else:
                self.status_text.set("Bild 'Themenlistenhelfer.png' nicht gefunden!")
        except Exception as e:
            self.status_text.set(f"Bildfehler: {e}")
    def excel_bearbeiten(self):
        path = self.excel_path.get()
        if not path or not os.path.exists(path):
            messagebox.showerror("Fehler", "Die Datei 'Auswahl Teilnehmende zu Lernbereichen.xlsx' wurde nicht gefunden!")
            return
        try:
            os.startfile(path)
        except Exception as e:
            messagebox.showerror("Fehler", f"Konnte Excel-Datei nicht öffnen: {e}")
    def batch_starten(self):
        path = self.excel_path.get()
        if not path or not os.path.exists(path):
            messagebox.showerror("Fehler", "Die Datei 'Auswahl Teilnehmende zu Lernbereichen.xlsx' wurde nicht gefunden!")
            return
        # Prüfen, ob Datensätze zu verarbeiten sind
        try:
            df = pd.read_excel(path, sheet_name='Teilnehmende')
            df = df[df['Verarbeiten'].str.lower() == 'ja']
            if df.empty:
                messagebox.showinfo("Hinweis", "In der Excel-Datei sind keine Datensätze aktiviert.")
                return
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Einlesen der Excel-Datei: {e}")
            return
        self.progress.set(0)
        self.progressbar.start(10)  # Animation starten
        threading.Thread(target=self.batch_ablauf, args=(path,), daemon=True).start()
    def batch_ablauf(self, path):
        pythoncom.CoInitialize()
        def set_progress(val):
            self.progress.set(val)
            self.root.update_idletasks()
            if val >= 100:
                self.progress_label.config(text="Fertig")
            else:
                self.progress_label.config(text="")
        def gui_callback(msg):
            # Zeige Status- oder Fehlermeldungen im Status-Label an
            self.status_text.set(msg)
            self.root.update_idletasks()
        try:
            set_progress(10)
            thema_result = erstelle_themenlisten(path, gui_callback)
            set_progress(50)
            email_result = erstelle_emails(path, gui_callback)
            set_progress(80)
            geloescht = loesche_alte_themenlisten(path, gui_callback)
            set_progress(100)
            if thema_result.get('aborted'):
                self.status_text.set("Abbruch: Es konnten keine gültigen Word-Vorlagen gefunden werden.")
            elif thema_result.get('created', 0) == 0:
                self.status_text.set("Keine Themenlisten erstellt. Bitte Vorlagen/Lernbereiche prüfen.")
            elif email_result.get('aborted'):
                self.status_text.set("Themenlisten erstellt, aber E-Mail-Erstellung wegen fehlender Vorlagen abgebrochen.")
            elif email_result.get('drafts', 0) == 0:
                self.status_text.set("Themenlisten erstellt, aber keine E-Mail-Entwürfe erzeugt (Vorlagen/Konfiguration/Anhänge prüfen).")
            else:
                self.status_text.set(f"Abgeschlossen: {thema_result.get('created', 0)} Themenlisten, {email_result.get('drafts', 0)} E-Mail-Entwürfe, {geloescht} Dateien bereinigt.")
        except Exception as e:
            set_progress(0)
            self.status_text.set(f"Fehler: {e}")
        finally:
            self.progressbar.stop()

if __name__ == "__main__":
    root = tk.Tk()
    # Icon setzen
    import os
    resource_dirs = get_resource_dirs()
    icon_path = resolve_path(
        *[os.path.join(d, "assets", "icons", "Themenlistenhelfer256.ico") for d in resource_dirs],
        *[os.path.join(d, "Themenlistenhelfer256.ico") for d in resource_dirs],
        must_exist=True
    )
    if icon_path and os.path.exists(icon_path):
        root.iconbitmap(icon_path)
    app = ThemenlistenApp(root)
    root.mainloop()
